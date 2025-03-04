from docx import Document

# Create a new Document for updated plan including entrepreneurship aspects
document = Document()

# Title and introduction
document.add_heading('PLAN DE ÁREA DE TECNOLOGÍA E INFORMÁTICA GRADO: ONCE', 0)
document.add_paragraph(
    'Este documento presenta el plan de contenidos para la enseñanza de JavaScript '
    'desde lo básico hasta lo avanzado, distribuido en cuatro períodos. Además, se incluyen aspectos de emprendimiento tecnológico, '
    'enfocados en la construcción de páginas web con JavaScript, para promover la innovación y el desarrollo de proyectos emprendedores.'
)

# ---------------------------
# PERÍODO 1: INTRODUCCIÓN A JAVASCRIPT Y FUNDAMENTOS BÁSICOS
document.add_heading('PERÍODO 1: INTRODUCCIÓN A JAVASCRIPT Y FUNDAMENTOS BÁSICOS', level=1)
document.add_paragraph(
    'Temática: Introducción a la programación en JavaScript, sintaxis básica, variables, tipos de datos, '
    'operadores, estructuras de control, funciones básicas y manipulación inicial del DOM.'
)
document.add_paragraph('Estándares:')
document.add_paragraph('● Conoce y utiliza la sintaxis básica y estructuras de control en JavaScript.')
document.add_paragraph('● Aplica buenas prácticas en la escritura de código.')
document.add_paragraph('● Desarrolla ejercicios básicos para la solución de problemas.')
document.add_paragraph('● Introduce fundamentos que servirán de base para proyectos tecnológicos y emprendimientos digitales.')

table1 = document.add_table(rows=2, cols=8)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'CONTENIDO TEMÁTICO'
hdr_cells[1].text = 'LOGRO'
hdr_cells[2].text = 'INDICADOR DE LOGRO'
hdr_cells[3].text = 'COMPETENCIA'
hdr_cells[4].text = 'ACTIVIDADES PEDAGÓGICAS'
hdr_cells[5].text = 'INDICADORES DE EVALUACIÓN'
hdr_cells[6].text = 'RECURSOS'
hdr_cells[7].text = 'METAS DE CALIDAD'

row_cells = table1.rows[1].cells
row_cells[0].text = ('Introducción a JavaScript, sintaxis, variables, tipos de datos, operadores, '
                       'estructuras de control, funciones básicas, manipulación del DOM.')
row_cells[1].text = 'Comprende la sintaxis básica de JavaScript y escribe scripts simples.'
row_cells[2].text = 'Crea programas simples utilizando estructuras de control.'
row_cells[3].text = 'Cognitiva, Razonamiento lógico'
row_cells[4].text = 'Ejercicios prácticos, mini proyectos y resolución de problemas en aula.'
row_cells[5].text = 'Evaluación de código y ejercicios resueltos.'
row_cells[6].text = 'Computador, navegador web, editor de código, documentación online.'
row_cells[7].text = 'El 80% de los estudiantes demuestran comprensión de los conceptos básicos.'

# ---------------------------
# PERÍODO 2: JAVASCRIPT INTERMEDIO: FUNCIONES AVANZADAS Y MANEJO DEL DOM
document.add_heading('PERÍODO 2: JAVASCRIPT INTERMEDIO: FUNCIONES AVANZADAS Y MANEJO DEL DOM', level=1)
document.add_paragraph(
    'Temática: Introducción a funciones avanzadas (arrow functions, callbacks), manejo de arreglos y objetos, '
    'manipulación del DOM y eventos. Se enfatiza la aplicación de estos conocimientos en el desarrollo de prototipos de páginas web con potencial emprendedor.'
)
document.add_paragraph('Estándares:')
document.add_paragraph('● Diseña y utiliza funciones avanzadas y estructuras de datos en JavaScript.')
document.add_paragraph('● Implementa interacciones básicas con el DOM y eventos.')
document.add_paragraph('● Desarrolla programas que apliquen programación funcional.')
document.add_paragraph('● Introduce conceptos que faciliten la creación de prototipos para emprendimientos digitales.')

table2 = document.add_table(rows=2, cols=8)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'CONTENIDO TEMÁTICO'
hdr_cells[1].text = 'LOGRO'
hdr_cells[2].text = 'INDICADOR DE LOGRO'
hdr_cells[3].text = 'COMPETENCIA'
hdr_cells[4].text = 'ACTIVIDADES PEDAGÓGICAS'
hdr_cells[5].text = 'INDICADORES DE EVALUACIÓN'
hdr_cells[6].text = 'RECURSOS'
hdr_cells[7].text = 'METAS DE CALIDAD'

row_cells = table2.rows[1].cells
row_cells[0].text = ('Funciones avanzadas (arrow functions, callbacks), manejo de arreglos y objetos, '
                       'manipulación del DOM y eventos.')
row_cells[1].text = 'Aplica conceptos intermedios de JavaScript en la manipulación del DOM y eventos.'
row_cells[2].text = 'Realiza proyectos que integren funciones avanzadas y manipulación del DOM.'
row_cells[3].text = 'Cognitiva, Analítica'
row_cells[4].text = 'Talleres prácticos, ejercicios de manipulación de elementos web y mini aplicaciones interactivas.'
row_cells[5].text = 'Evaluación de aplicaciones web y revisión de código.'
row_cells[6].text = 'Computador, navegador, editor de código, documentación (MDN).'
row_cells[7].text = 'El 75% de los estudiantes logran manipular el DOM de forma efectiva.'

# ---------------------------
# PERÍODO 3: JAVASCRIPT AVANZADO: ASINCRONÍA, APIs Y PROGRAMACIÓN ORIENTADA A OBJETOS
document.add_heading('PERÍODO 3: JAVASCRIPT AVANZADO: ASINCRONÍA, APIs Y PROGRAMACIÓN ORIENTADA A OBJETOS', level=1)
document.add_paragraph(
    'Temática: Programación asincrónica con Promises y async/await, consumo de APIs con fetch, '
    'y fundamentos de la programación orientada a objetos en JavaScript. Se busca que los estudiantes desarrollen soluciones web robustas que puedan ser la base para proyectos emprendedores.'
)
document.add_paragraph('Estándares:')
document.add_paragraph('● Comprende y aplica técnicas de programación asincrónica en JavaScript.')
document.add_paragraph('● Conecta aplicaciones JavaScript con servicios web mediante APIs.')
document.add_paragraph('● Aplica principios de programación orientada a objetos.')
document.add_paragraph('● Fomenta la integración de estos conocimientos en proyectos con potencial de emprendimiento.')

table3 = document.add_table(rows=2, cols=8)
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'CONTENIDO TEMÁTICO'
hdr_cells[1].text = 'LOGRO'
hdr_cells[2].text = 'INDICADOR DE LOGRO'
hdr_cells[3].text = 'COMPETENCIA'
hdr_cells[4].text = 'ACTIVIDADES PEDAGÓGICAS'
hdr_cells[5].text = 'INDICADORES DE EVALUACIÓN'
hdr_cells[6].text = 'RECURSOS'
hdr_cells[7].text = 'METAS DE CALIDAD'

row_cells = table3.rows[1].cells
row_cells[0].text = ('Promises, async/await, consumo de APIs (fetch) y fundamentos de la programación orientada a objetos.')
row_cells[1].text = 'Desarrolla aplicaciones que integran consumo de APIs y POO.'
row_cells[2].text = 'Implementa código asincrónico y estructuras de POO en proyectos.'
row_cells[3].text = 'Cognitiva, Resolución de problemas, Creativa'
row_cells[4].text = 'Proyectos prácticos, desarrollo de módulos, integración de APIs y ejercicios colaborativos.'
row_cells[5].text = 'Evaluación de proyectos, pruebas prácticas y análisis de código.'
row_cells[6].text = 'Computador, navegador, editores de código, documentación (MDN, API docs).'
row_cells[7].text = 'El 70% de los estudiantes desarrollan aplicaciones que consumen APIs y aplican POO.'

# ---------------------------
# PERÍODO 4: DESARROLLO DE PROYECTOS WEB INTERACTIVOS CON JAVASCRIPT
document.add_heading('PERÍODO 4: DESARROLLO DE PROYECTOS WEB INTERACTIVOS CON JAVASCRIPT', level=1)
document.add_paragraph(
    'Temática: Desarrollo integral de una aplicación web interactiva utilizando JavaScript, '
    'integrando los conocimientos adquiridos en los períodos anteriores. Este período enfatiza la realización de proyectos con un enfoque práctico que puede transformarse en iniciativas emprendedoras.'
)
document.add_paragraph('Estándares:')
document.add_paragraph('● Aplica de manera integral los conocimientos de JavaScript en el desarrollo de proyectos.')
document.add_paragraph('● Diseña y programa aplicaciones web interactivas con buenas prácticas.')
document.add_paragraph('● Desarrolla un proyecto final que demuestre el dominio de JavaScript.')
document.add_paragraph('● Integra conceptos de innovación y emprendimiento tecnológico en el desarrollo de proyectos web.')

table4 = document.add_table(rows=2, cols=8)
hdr_cells = table4.rows[0].cells
hdr_cells[0].text = 'CONTENIDO TEMÁTICO'
hdr_cells[1].text = 'LOGRO'
hdr_cells[2].text = 'INDICADOR DE LOGRO'
hdr_cells[3].text = 'COMPETENCIA'
hdr_cells[4].text = 'ACTIVIDADES PEDAGÓGICAS'
hdr_cells[5].text = 'INDICADORES DE EVALUACIÓN'
hdr_cells[6].text = 'RECURSOS'
hdr_cells[7].text = 'METAS DE CALIDAD'

row_cells = table4.rows[1].cells
row_cells[0].text = ('Diseño y desarrollo de una aplicación web interactiva, integración de APIs, manejo avanzado del DOM, '
                       'depuración y optimización.')
row_cells[1].text = 'Desarrolla un proyecto final completo utilizando JavaScript.'
row_cells[2].text = ('Presenta un proyecto funcional con interactividad, consumo de APIs y aplicación de buenas prácticas en el código.')
row_cells[3].text = 'Cognitiva, Trabajo en equipo, Creativa, Comunicación'
row_cells[4].text = ('Desarrollo de proyecto en equipo o individual, presentaciones, revisión de código y retroalimentación.')
row_cells[5].text = 'Evaluación del proyecto final, presentaciones y análisis de código.'
row_cells[6].text = 'Computador, navegador, editor de código, herramientas de desarrollo, servidores web.'
row_cells[7].text = 'El 70% de los estudiantes presentan proyectos web interactivos completos y funcionales.'

# ---------------------------
# NUEVA SECCIÓN: ASPECTOS DE EMPRENDIMIENTO TECNOLÓGICO EN EL DESARROLLO WEB CON JAVASCRIPT
document.add_heading('ASPECTOS DE EMPRENDIMIENTO TECNOLÓGICO EN EL DESARROLLO WEB CON JAVASCRIPT', level=1)
document.add_paragraph(
    'Esta sección aborda la integración de competencias emprendedoras en el ámbito del desarrollo web, '
    'focalizándose en la construcción de páginas web con JavaScript como herramienta para innovar y generar proyectos con viabilidad de negocio. '
    'Se incluyen temas relacionados con la identificación de oportunidades de mercado, diseño de modelos de negocio, estrategias de marketing digital, '
    'creación de prototipos (MVP), y gestión de proyectos tecnológicos.'
)
document.add_paragraph('Estándares:')
document.add_paragraph('● Identifica oportunidades de negocio en el entorno digital.')
document.add_paragraph('● Diseña modelos de negocio y estrategias de marketing aplicables a proyectos web.')
document.add_paragraph('● Desarrolla prototipos y MVPs que integren funcionalidades de JavaScript para validar ideas emprendedoras.')
document.add_paragraph('● Aplica metodologías ágiles y de gestión de proyectos en iniciativas tecnológicas.')

table5 = document.add_table(rows=2, cols=8)
hdr_cells = table5.rows[0].cells
hdr_cells[0].text = 'CONTENIDO TEMÁTICO'
hdr_cells[1].text = 'LOGRO'
hdr_cells[2].text = 'INDICADOR DE LOGRO'
hdr_cells[3].text = 'COMPETENCIA'
hdr_cells[4].text = 'ACTIVIDADES PEDAGÓGICAS'
hdr_cells[5].text = 'INDICADORES DE EVALUACIÓN'
hdr_cells[6].text = 'RECURSOS'
hdr_cells[7].text = 'METAS DE CALIDAD'

row_cells = table5.rows[1].cells
row_cells[0].text = ('Análisis de oportunidades de negocio digital, diseño de modelos de negocio, estrategias de marketing digital, '
                       'creación de prototipos (MVP) y gestión de proyectos tecnológicos aplicados a el desarrollo de páginas web.')
row_cells[1].text = 'Desarrolla una propuesta de emprendimiento tecnológico basada en una aplicación web interactiva.'
row_cells[2].text = 'Presenta un proyecto o prototipo que integre funcionalidades de JavaScript con un modelo de negocio innovador.'
row_cells[3].text = 'Emprendimiento, Creatividad, Gestión de proyectos'
row_cells[4].text = ('Talleres de ideación, análisis de casos de éxito, elaboración de plan de negocio, desarrollo de prototipos y presentaciones de proyectos.')
row_cells[5].text = 'Evaluación de propuestas de negocio, viabilidad del prototipo y claridad en la estrategia emprendedora.'
row_cells[6].text = 'Computador, herramientas de diseño, software de prototipado, internet y documentación de marketing digital.'
row_cells[7].text = 'El 65% de los estudiantes presentan propuestas de emprendimiento tecnológico viables y bien estructuradas.'

# ---------------------------
# Resumen por Períodos
document.add_heading('Resumen por Períodos', level=1)
table_summary = document.add_table(rows=6, cols=3)
hdr_cells = table_summary.rows[0].cells
hdr_cells[0].text = 'Período'
hdr_cells[1].text = 'Tema Principal'
hdr_cells[2].text = 'Estándar Principal'

row_cells = table_summary.rows[1].cells
row_cells[0].text = '1'
row_cells[1].text = 'Fundamentos de JavaScript'
row_cells[2].text = 'Comprensión de sintaxis y estructuras básicas'

row_cells = table_summary.rows[2].cells
row_cells[0].text = '2'
row_cells[1].text = 'JavaScript Intermedio'
row_cells[2].text = 'Manipulación del DOM y funciones avanzadas'

row_cells = table_summary.rows[3].cells
row_cells[0].text = '3'
row_cells[1].text = 'JavaScript Avanzado'
row_cells[2].text = 'Programación asincrónica, APIs y POO'

row_cells = table_summary.rows[4].cells
row_cells[0].text = '4'
row_cells[1].text = 'Proyecto Final'
row_cells[2].text = 'Desarrollo integral de una aplicación web interactiva'

row_cells = table_summary.rows[5].cells
row_cells[0].text = 'Emprendimiento'
row_cells[1].text = 'Emprendimiento Tecnológico'
row_cells[2].text = 'Integración de competencias emprendedoras en proyectos web'

# Save the document
document.save('Plan_JS_Emprendimiento.docx')
print("El archivo 'Plan_JS_Emprendimiento.docx' ha sido generado exitosamente.")
