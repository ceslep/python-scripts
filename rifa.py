from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE

# Crear el documento
document = Document()

# Configurar tamaño de página y márgenes para tamaño carta (8.5 x 11 pulgadas)
section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Agregar el encabezado centrado
header_paragraph = document.add_paragraph()
header_paragraph.alignment = 1  # 0=izquierda, 1=centrado, 2=derecha
header_run = header_paragraph.add_run(
    "Los estudiantes de grado 1102 del Instituto Guática hacen una rifa con el fin de recaudar fondos para las actividades de finalización de estudios"
)
header_run.bold = True
header_run.font.size = Pt(14)

# Agregar información del premio centrada
prize_paragraph = document.add_paragraph()
prize_paragraph.alignment = 1
prize_run = prize_paragraph.add_run("Premio: Combo de Pollo frito Frisby")
prize_run.bold = True
prize_run.font.size = Pt(12)

# Agregar un párrafo en blanco para separar
document.add_paragraph()

# ---------------------------------------------------------------------------
# Calcular la altura de cada fila para que la tabla ocupe el espacio restante.
# Se asume:
# - Página: 11" de alto
# - Márgenes superior e inferior: 0.5" cada uno  --> espacio vertical total: 11 - 1 = 10"
# - Espacio usado por los párrafos de encabezado y premio (aproximado): 1.5"
# Entonces, espacio disponible para la tabla: 10 - 1.5 = 8.5" 
#
# La tabla tendrá 21 filas (1 fila de encabezado y 20 para participantes).
# Altura por fila = 8.5" / 21 ≈ 0.405"
# ---------------------------------------------------------------------------
row_height = Inches(8.5 / 21)

# Crear la tabla: 21 filas y 3 columnas
table = document.add_table(rows=21, cols=3)
table.style = 'Table Grid'

# Rellenar la primera fila (encabezados de la tabla)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'N°'
hdr_cells[1].text = 'Nombre'
hdr_cells[2].text = 'Teléfono'

# Configurar la altura de cada fila de la tabla de forma fija (exacta)
for row in table.rows:
    row.height = row_height
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# Rellenar la primera columna de las filas de participantes (filas 2 a 21)
for i in range(1, 21):
    row_cells = table.rows[i].cells
    row_cells[0].text = str(i)
    # Las celdas "Nombre" y "Teléfono" se dejan en blanco para completarlas manualmente

# Guardar el documento
document.save("rifa.docx")
